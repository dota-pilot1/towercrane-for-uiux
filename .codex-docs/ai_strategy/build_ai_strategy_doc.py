from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("/Users/terecal/Desktop/이력서/ai 전략팀 지원 전략.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(18, 31, 45)
MUTED = RGBColor(87, 96, 111)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
BORDER = "D9E2EC"
WHITE = RGBColor(255, 255, 255)
FONT = "Malgun Gothic"


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def paragraph_border_bottom(paragraph, color="2E74B5", size="10"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def set_fixed_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120, col_widths=None):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    if col_widths:
        grid = table._tbl.tblGrid
        if grid is None:
            grid = OxmlElement("w:tblGrid")
            table._tbl.insert(0, grid)
        for child in list(grid):
            grid.remove(child)
        for width in col_widths:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(width))
            grid.append(col)
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                set_fixed_cell_width(cell, col_widths[idx])


def add_body_paragraph(doc, text="", bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, 11, INK, True)
        r2 = p.add_run(text[len(bold_prefix) :])
        set_run_font(r2, 11, INK)
    else:
        r = p.add_run(text)
        set_run_font(r, 11, INK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, 16, BLUE, True)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_run_font(r, 13, BLUE, True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        set_run_font(r, 12, DARK_BLUE, True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run_font(r, 11, INK)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, 9360, 120, [9360])
    set_table_borders(table, "D9E2EC", "6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, 11, DARK_BLUE, True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.10
    r2 = p2.add_run(body)
    set_run_font(r2, 10.5, INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)
    return table


def add_matrix_table(doc, headers, rows, widths, bold_first_col=True):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_width(table, 9360, 120, widths)
    set_table_borders(table, "D9E2EC", "4")
    for idx, text in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, 10.5, DARK_BLUE, True)

    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cell = cells[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.10
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run_font(r, 10.0, INK, bold=(bold_first_col and idx == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_strategy_table(doc):
    rows = [
        ("공통 프로젝트 자산화", "프로젝트 템플릿, 공통 컴포넌트/모듈, API 문서, 테스트 도구, 배포 정보를 재사용 가능한 자산으로 축적"),
        ("프로토타입 모음 구축", "신규 기능 아이디어와 실험 결과를 모아 빠른 의사결정과 후속 개발의 출발점으로 활용"),
        ("반복 개발 자동화", "문서 생성, 코드 리뷰, 테스트 실행, 배포 체크리스트 등 반복 업무를 AI와 스크립트로 자동화"),
        ("온보딩·실습 환경", "신입/주니어 개발자가 표준 개발 방식, SQL, API, 제품 기획 흐름을 직접 실습할 수 있는 환경 제공"),
        ("표준 개발 방식 정립", "기획-설계-구현-리뷰-보안 점검을 Golden Path로 정의해 프로젝트 품질 편차를 줄임"),
    ]
    add_matrix_table(doc, ["전략 방향", "구현 내용"], rows, [2450, 6910])


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name in ["Heading 1", "Heading 2", "Heading 3", "List Bullet"]:
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    r = header.add_run("AI Strategy Brief")
    set_run_font(r, 9, MUTED, True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    r = footer.add_run("AI 전략팀 지원 전략")
    set_run_font(r, 9, MUTED)


def add_masthead(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("AI 전략팀 지원 전략")
    set_run_font(r, 23, INK, True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("사내 개발 자산 플랫폼과 AI 활용 체계 구축 제안")
    set_run_font(r, 13.5, MUTED)

    metadata = [
        ("문서 목적:", "AI 전략팀 지원용 핵심 제안 요약"),
        ("작성일:", "2026.05.15"),
        ("핵심 메시지:", "AI 전략은 챗봇 도입을 넘어, 사내 개발 자산을 축적·재사용하는 플랫폼 전략으로 확장되어야 합니다."),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label + " ")
        set_run_font(r, 10.5, INK, True)
        r2 = p.add_run(value)
        set_run_font(r2, 10.5, INK)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(8)
    paragraph_border_bottom(rule, "2E74B5", "12")


def build():
    doc = Document()
    configure_document(doc)
    add_masthead(doc)

    add_callout(
        doc,
        "한 줄 요약",
        "AI 시대의 경쟁력은 개별 도구 사용 능력이 아니라, 회사 내부의 개발 자산·지식·표준 업무 방식을 플랫폼화해 반복 가능한 생산성으로 만드는 데 있습니다.",
    )

    add_heading(doc, "1. 제안 배경", 1)
    add_body_paragraph(
        doc,
        "AI 도구를 단순히 업무 보조 수단으로 도입하면 개인별 활용 편차가 커지고, 결과물이 조직 자산으로 남기 어렵습니다. 따라서 AI 전략팀은 도구 도입보다 먼저 사내 개발 자산을 체계적으로 축적하고 재사용할 수 있는 구조를 만드는 방향으로 접근해야 합니다.",
    )
    add_bullet(doc, "공통 프로젝트 자산화와 표준 개발 방식 정립")
    add_bullet(doc, "프로토타입, 업무 노트, 개발 가이드, API 문서의 중앙화")
    add_bullet(doc, "반복 개발 업무 자동화와 신입·주니어 개발자 온보딩 체계화")
    add_bullet(doc, "인재 양성용 실습 환경과 AI 활용 교육 기반 제공")

    add_heading(doc, "2. 참고 모델: Spotify Backstage", 1)
    add_body_paragraph(
        doc,
        "Spotify Backstage는 내부 개발자 포털(IDP)의 대표 모델입니다. 개발자가 매번 처음부터 환경을 구성하지 않고, 회사가 정한 표준 방식에 따라 빠르게 개발을 시작하도록 돕는 시스템입니다.",
    )
    add_matrix_table(
        doc,
        ["Backstage 기능", "AI 전략팀 관점의 의미"],
        [
            ("Software Catalog", "서비스, API, 라이브러리, 담당자, 문서, 배포 정보를 한곳에서 관리"),
            ("Software Templates", "표준 템플릿 기반으로 신규 프로젝트를 빠르게 생성"),
            ("Internal Developer Portal", "배포, 문서, 로그, 모니터링, 권한 요청 등 개발자 업무 진입점을 통합"),
            ("Golden Path", "회사가 권장하는 표준 개발·배포 방식을 제공해 품질 편차를 줄임"),
        ],
        [2600, 6760],
    )

    add_heading(doc, "3. 국내 유사 사례", 1)
    add_body_paragraph(
        doc,
        "국내 기업들이 Backstage를 직접 사용한다고 단정하기보다는, 같은 문제를 자체 플랫폼과 플랫폼 엔지니어링 방식으로 해결하고 있다고 보는 것이 적절합니다.",
    )
    add_matrix_table(
        doc,
        ["사례", "참고 포인트"],
        [
            ("토스", "개발자 생산성, 배포 파이프라인, 공통 모듈, 내부 도구를 적극적으로 구축"),
            ("카카오페이", "WECAN 개발자 플랫폼 및 내부 개발자 포털 방향의 조직 운영"),
            ("카카오뱅크", "금융권 환경에 맞춘 클라우드/플랫폼 엔지니어링 체계 운영"),
        ],
        [2300, 7060],
    )

    add_heading(doc, "4. 제안 방향: 사내 개발 자산 플랫폼", 1)
    add_strategy_table(doc)
    add_body_paragraph(
        doc,
        "이 플랫폼은 단순 문서 저장소가 아니라 프로젝트 템플릿, 공통 모듈, API 문서, SQL 실습, 업무 노트, 프로토타입, 온보딩 과정, AI 기반 코드 리뷰와 문서 생성을 연결하는 실행 기반이 되어야 합니다.",
    )

    add_heading(doc, "5. 구축 경험과 확장 계획", 1)
    add_matrix_table(
        doc,
        ["구분", "내용", "확장 방향"],
        [
            ("구축한 기본 플랫폼", "https://hibot-docu.com/", "AI 전략팀의 내부 자산 플랫폼 PoC로 확장 가능"),
            ("현재 기능", "업무 관리, SQL 연습, 프로토타입, 챌린지, 회의실, 스터디 다이어리", "개발자 실습·업무 기록·프로젝트 공유 공간으로 고도화"),
            ("업무 보상 시스템", "Stack Overflow for Teams/Internal 방식의 사내 지식 공유 플랫폼", "질문, 답변, 문서화, 문제 해결 기여를 축적하고 보상 지표로 활용"),
            ("AI 지식 챗봇", "축적된 질문·답변·문서를 AI와 연결", "사내 지식 검색, 온보딩 Q&A, 장애 대응 지원으로 확장"),
        ],
        [1900, 3450, 4010],
    )

    add_heading(doc, "6. AI 활용 기술 전파", 1)
    add_body_paragraph(
        doc,
        "AI 활용의 핵심은 프롬프트를 잘 쓰는 것에 그치지 않고, AI가 일하는 방식을 통제·검증·개선하는 개발 프로세스를 조직에 전파하는 것입니다.",
    )
    add_matrix_table(
        doc,
        ["개념/도구", "핵심 내용"],
        [
            ("하네스 엔지니어링", "AI가 멋대로 행동하지 않도록 기획-설계-구현-리뷰 단계를 구조화해 결과물 품질을 통제"),
            ("컴파운드 엔지니어링", "실행·검토 과정의 실수를 기록하고 다음 단계에 반영해 같은 실수를 줄이는 방식"),
            ("Codex Desktop / GPT-5.5", "CLI보다 접근성이 높은 데스크톱 환경과 고성능 모델을 활용해 개발·검토 속도 향상"),
            ("Superpowers", "초기 아이디어를 요구사항, 기획, 브레인스토밍 결과로 정리하는 개발 보조 스킬로 활용"),
            ("G-Stack / CSO", "CEO, 보안 책임자 등 역할 기반 검토와 보안 체크리스트 기반 점검에 활용"),
        ],
        [2600, 6760],
    )
    add_heading(doc, "AI 기반 개발 프로세스", 2)
    add_matrix_table(
        doc,
        ["단계", "운영 방식"],
        [
            ("1. 기획", "사용자의 대략적 아이디어를 AI가 요구사항, 기능 목록, 우선순위로 구체화"),
            ("2. 구현 계획", "기술 스택과 구조를 정하고 설계 문서/Spec을 작성"),
            ("3. 구현 및 검토", "서브 에이전트가 구현하고 메인 에이전트가 리뷰·수정"),
            ("4. 보안 점검", "CSO 관점에서 개인정보, 로컬 저장, 권한, 배포 위험을 체크"),
            ("5. 회고 축적", "실수와 개선점을 기록해 다음 프로젝트의 하네스와 체크리스트에 반영"),
        ],
        [2100, 7260],
    )
    add_callout(
        doc,
        "운영 원칙",
        "기획 문서(Spec)를 프로젝트의 Single Source of Truth로 두고, AI가 바꾸는 코드와 기준 문서를 분리해 관리합니다. 보안 점검과 고성능 모델 사용에는 토큰을 아끼지 않고 품질 기준을 높입니다.",
    )

    add_heading(doc, "7. 기술 도입 및 업무 적용", 1)
    add_matrix_table(
        doc,
        ["영역", "적용 방향"],
        [
            ("Spring AI", "사내 서비스, 지식 검색, 문서 생성, 코드 리뷰 기능을 백엔드 서비스와 연결"),
            ("Tauri / Flutter", "데스크톱 앱, 모바일 앱, 내부 도구를 빠르게 프로토타입으로 제작"),
            ("스타트업 방식", "작게 만들고 빠르게 검증하는 실험 문화를 도입"),
            ("토스뱅크식 장점", "금융권 수준의 안정성, 자동화, 플랫폼 엔지니어링 관점을 참고"),
            ("채용 시스템", "지원자 과제 리뷰, 면접 질문 생성, 평가 기록 정리, 온보딩 연결에 AI 적용"),
            ("업무 관리", "회의록, 업무 노트, 진행 상황 요약, 리스크 정리, 후속 액션 자동화"),
            ("제품 개발", "기획서, 요구사항, 프로토타입, 코드 리뷰, 테스트 케이스, 릴리즈 문서 생성에 AI 적용"),
        ],
        [2300, 7060],
    )

    add_heading(doc, "8. 기대 효과", 1)
    add_bullet(doc, "반복 개발 시간 단축과 프로젝트 착수 속도 향상")
    add_bullet(doc, "신규 인력 온보딩 및 주니어 개발자 성장 체계 강화")
    add_bullet(doc, "공통 모듈, 문서, 템플릿, 지식 답변의 사내 자산화")
    add_bullet(doc, "프로젝트 품질 표준화와 개발자 생산성 향상")
    add_bullet(doc, "채용, 업무 관리, 제품 개발 과정에서 AI 활용 범위 확대")
    add_bullet(doc, "개인 중심 AI 활용을 조직 차원의 생산성 시스템으로 전환")

    add_callout(
        doc,
        "지원 포지셔닝",
        "이미 구축한 기본 플랫폼을 바탕으로 사내 개발자 포털, 업무 보상 시스템, 지식 기반 챗봇, AI 개발 프로세스 교육을 연결해 AI 전략팀의 실행형 플랫폼 과제를 빠르게 추진할 수 있습니다.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
